"""
Estrattore .docx via python-docx.

Strategia:
  - Headings (Heading 1/2/3...) demarcano nuove Section
  - Paragrafi vanno nella sezione corrente
  - Le tabelle vengono serializzate riga per riga (come CSV inline)
"""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document  # python-docx

from core.dataset.extracted import ExtractedDocument, Section
from core.dataset.extractors.base import Extractor, ExtractorError

logger = logging.getLogger(__name__)


class DocxExtractor(Extractor):
    """Estrattore per file Word .docx."""

    SUPPORTED_EXTENSIONS = (".docx",)

    def extract(self, path: Path) -> ExtractedDocument:
        self._validate_file(path)

        try:
            doc = Document(str(path))
        except Exception as exc:  # noqa: BLE001
            raise ExtractorError(f"DOCX non leggibile {path}: {exc}") from exc

        sections: list[Section] = []
        current_title: str | None = None
        current_buffer: list[str] = []
        all_text: list[str] = []

        def flush_section() -> None:
            """Salva la sezione corrente e resetta il buffer."""
            if current_buffer:
                text = "\n\n".join(current_buffer).strip()
                if text:
                    sections.append(
                        Section(
                            title=current_title,
                            text=text,
                            metadata={},
                        )
                    )
                current_buffer.clear()

        # Iteriamo sui paragrafi nell'ordine del documento
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style = (para.style.name or "").lower() if para.style else ""

            if style.startswith("heading"):
                # Nuova sezione: chiudiamo la precedente
                flush_section()
                current_title = text
            else:
                current_buffer.append(text)

            all_text.append(text)

        # Tabelle → serializzate dopo il body, in una sezione "Tabelle"
        if doc.tables:
            tables_text_parts: list[str] = []
            for t_idx, table in enumerate(doc.tables, start=1):
                rows_text = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows_text.append(" | ".join(cells))
                table_text = f"[Tabella {t_idx}]\n" + "\n".join(rows_text)
                tables_text_parts.append(table_text)
                all_text.append(table_text)

            # Aggiungi al buffer corrente o crea sezione dedicata
            current_buffer.append("\n\n".join(tables_text_parts))

        # Flush finale
        flush_section()

        full_text = "\n\n".join(all_text)

        return ExtractedDocument(
            text=full_text,
            source_format="docx",
            sections=sections,
            metadata={
                "section_count": len(sections),
                "table_count": len(doc.tables),
                "paragraph_count": len(doc.paragraphs),
                "file_size_bytes": path.stat().st_size,
            },
        )
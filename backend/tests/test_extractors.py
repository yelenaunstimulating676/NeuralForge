"""
Test per gli estrattori del Dataset Engine.

Strategia: ogni test crea un file temporaneo in tmp_path con contenuto
controllato, lo passa all'estrattore, verifica il risultato.
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest

from core.dataset.extracted import ExtractedDocument
from core.dataset.extractors import (
    UnsupportedFormatError,
    extract_file,
    get_extractor_for_path,
    is_supported_extension,
    supported_extensions,
)
from core.dataset.extractors.base import ExtractorError


# ---------------------------------------------------------------------------
# Router
# ---------------------------------------------------------------------------


class TestRouter:
    def test_supported_extensions_set(self):
        ext = supported_extensions()
        assert ".pdf" in ext
        assert ".txt" in ext
        assert ".csv" in ext
        assert ".tsv" in ext
        assert ".json" in ext
        assert ".jsonl" in ext
        assert ".docx" in ext
        assert ".md" in ext

    def test_is_supported_extension(self):
        assert is_supported_extension(Path("file.pdf")) is True
        assert is_supported_extension(Path("file.PDF")) is True  # case-insensitive
        assert is_supported_extension(Path("file.xyz")) is False

    def test_get_extractor_for_unsupported_raises(self, tmp_path):
        f = tmp_path / "x.xyz"
        f.write_text("data")
        with pytest.raises(UnsupportedFormatError):
            get_extractor_for_path(f)


# ---------------------------------------------------------------------------
# TXT
# ---------------------------------------------------------------------------


class TestTxtExtractor:
    def test_basic_utf8(self, tmp_path):
        f = tmp_path / "sample.txt"
        f.write_text("Hello world\nSecond line", encoding="utf-8")
        doc = extract_file(f)

        assert isinstance(doc, ExtractedDocument)
        assert doc.source_format == "txt"
        assert "Hello world" in doc.text
        assert "Second line" in doc.text
        assert doc.section_count == 1
        assert doc.metadata["encoding"] in {"utf-8", "ascii", "ASCII"}

    def test_md_works(self, tmp_path):
        f = tmp_path / "doc.md"
        f.write_text("# Title\n\nParagraph", encoding="utf-8")
        doc = extract_file(f)
        assert doc.source_format == "md"
        assert "# Title" in doc.text

    def test_empty_file_raises(self, tmp_path):
        f = tmp_path / "empty.txt"
        f.write_bytes(b"")
        with pytest.raises(ExtractorError, match="vuoto"):
            extract_file(f)

    def test_nonexistent_raises(self, tmp_path):
        with pytest.raises(ExtractorError, match="non trovato"):
            extract_file(tmp_path / "nope.txt")


# ---------------------------------------------------------------------------
# CSV
# ---------------------------------------------------------------------------


class TestCsvExtractor:
    def test_basic_csv(self, tmp_path):
        f = tmp_path / "data.csv"
        f.write_text(
            "name,age,city\nAlice,30,Roma\nBob,25,Milano\n",
            encoding="utf-8",
        )
        doc = extract_file(f)

        assert doc.source_format == "csv"
        assert doc.metadata["row_count"] == 2
        assert doc.metadata["column_count"] == 3
        assert doc.metadata["delimiter"] == ","
        assert doc.section_count == 2
        assert "Alice" in doc.sections[0].text
        assert "name: Alice" in doc.sections[0].text

    def test_tsv_with_tab(self, tmp_path):
        f = tmp_path / "data.tsv"
        f.write_text("a\tb\nx\ty\n", encoding="utf-8")
        doc = extract_file(f)
        assert doc.source_format == "tsv"
        assert doc.metadata["delimiter"] == "\t"
        assert doc.metadata["row_count"] == 1

    def test_semicolon_delimiter_detected(self, tmp_path):
        f = tmp_path / "euro.csv"
        f.write_text("col1;col2\n1;2\n3;4\n", encoding="utf-8")
        doc = extract_file(f)
        assert doc.metadata["delimiter"] == ";"
        assert doc.metadata["row_count"] == 2


# ---------------------------------------------------------------------------
# JSON / JSONL
# ---------------------------------------------------------------------------


class TestJsonExtractor:
    def test_array_of_objects(self, tmp_path):
        data = [
            {"q": "ciao?", "a": "ciao!"},
            {"q": "come va?", "a": "bene"},
        ]
        f = tmp_path / "data.json"
        f.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")

        doc = extract_file(f)
        assert doc.source_format == "json"
        assert doc.metadata["item_count"] == 2
        assert doc.section_count == 2
        assert "q: ciao?" in doc.sections[0].text
        assert "a: ciao!" in doc.sections[0].text

    def test_single_object(self, tmp_path):
        f = tmp_path / "obj.json"
        f.write_text('{"k": "v"}', encoding="utf-8")
        doc = extract_file(f)
        assert doc.metadata["item_count"] == 1
        assert "k: v" in doc.text

    def test_jsonl(self, tmp_path):
        f = tmp_path / "data.jsonl"
        f.write_text(
            '{"x": 1}\n{"x": 2}\n{"x": 3}\n', encoding="utf-8"
        )
        doc = extract_file(f)
        assert doc.source_format == "jsonl"
        assert doc.metadata["item_count"] == 3
        assert doc.metadata["is_jsonl"] is True

    def test_jsonl_with_invalid_line(self, tmp_path):
        f = tmp_path / "broken.jsonl"
        f.write_text(
            '{"ok": true}\nNOT JSON\n{"ok": false}\n', encoding="utf-8"
        )
        doc = extract_file(f)
        # 2 oggetti validi, warning per quello rotto
        assert doc.metadata["item_count"] == 2
        assert any("Riga 2" in w for w in doc.metadata["warnings"])

    def test_malformed_json_raises(self, tmp_path):
        f = tmp_path / "bad.json"
        f.write_text('{"unterminated":', encoding="utf-8")
        with pytest.raises(ExtractorError, match="malformato"):
            extract_file(f)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


class TestDocxExtractor:
    """
    python-docx richiede file binari Word veri. Generiamo al volo.
    """

    def test_basic_docx(self, tmp_path):
        from docx import Document

        f = tmp_path / "sample.docx"
        d = Document()
        d.add_heading("Capitolo 1", level=1)
        d.add_paragraph("Primo paragrafo.")
        d.add_paragraph("Secondo paragrafo.")
        d.add_heading("Capitolo 2", level=1)
        d.add_paragraph("Altro contenuto.")
        d.save(str(f))

        doc = extract_file(f)
        assert doc.source_format == "docx"
        assert "Primo paragrafo." in doc.text
        assert "Capitolo 1" in doc.text
        assert doc.section_count == 2  # due heading
        assert doc.sections[0].title == "Capitolo 1"
        assert doc.sections[1].title == "Capitolo 2"

    def test_docx_with_table(self, tmp_path):
        from docx import Document

        f = tmp_path / "tab.docx"
        d = Document()
        d.add_paragraph("Testo prima della tabella")
        table = d.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "2"
        d.save(str(f))

        doc = extract_file(f)
        assert doc.metadata["table_count"] == 1
        assert "A | B" in doc.text
        assert "1 | 2" in doc.text


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


class TestPdfExtractor:
    def test_basic_pdf(self, tmp_path):
        import fitz

        f = tmp_path / "sample.pdf"
        doc_pdf = fitz.open()
        page = doc_pdf.new_page()
        page.insert_text(
            (50, 100), "Hello PDF world", fontsize=14
        )
        page2 = doc_pdf.new_page()
        page2.insert_text((50, 100), "Pagina due", fontsize=14)
        doc_pdf.save(str(f))
        doc_pdf.close()

        doc = extract_file(f)
        assert doc.source_format == "pdf"
        assert doc.metadata["page_count"] == 2
        assert doc.section_count == 2
        assert "Hello PDF world" in doc.text
        assert "Pagina due" in doc.text
        assert doc.sections[0].metadata["page_number"] == 1

    def test_empty_pdf_warns(self, tmp_path):
        import fitz

        f = tmp_path / "empty.pdf"
        doc_pdf = fitz.open()
        doc_pdf.new_page()  # pagina senza testo
        doc_pdf.save(str(f))
        doc_pdf.close()

        doc = extract_file(f)
        assert any(
            "scansionato" in w.lower() for w in doc.metadata["warnings"]
        )


# ---------------------------------------------------------------------------
# ExtractedDocument helpers
# ---------------------------------------------------------------------------


class TestExtractedDocument:
    def test_to_dict_serializable(self, tmp_path):
        f = tmp_path / "x.txt"
        f.write_text("ciao", encoding="utf-8")
        doc = extract_file(f)
        d = doc.to_dict()

        # Serializzabile in JSON
        assert json.dumps(d)
        assert d["source_format"] == "txt"
        assert d["char_count"] == 4
        assert d["section_count"] == 1